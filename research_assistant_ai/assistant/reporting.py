from __future__ import annotations
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, Any, Optional
import pandas as pd

@dataclass
class ReportArtifacts:
    excel_path: Path
    word_path: Path
    figure_path: Optional[Path] = None

def export_excel(results: Dict[str, pd.DataFrame], out_path: Path) -> Path:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    with pd.ExcelWriter(out_path, engine="openpyxl") as writer:
        for name, df in results.items():
            safe = name[:31]
            df.to_excel(writer, sheet_name=safe, index=False)
    return out_path

def export_word(summary: Dict[str, Any], out_path: Path) -> Path:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    from docx import Document

    doc = Document()
    doc.add_heading("ResearchAssistantAI Run Summary", level=1)

    for section, payload in summary.items():
        doc.add_heading(str(section), level=2)
        if isinstance(payload, dict):
            for k, v in payload.items():
                doc.add_paragraph(f"{k}: {v}")
        else:
            doc.add_paragraph(str(payload))

    doc.save(out_path)
    return out_path

def plot_top_specs(spec_df: pd.DataFrame, out_path: Path, top_k: int = 10) -> Optional[Path]:
    if spec_df is None or spec_df.empty:
        return None
    out_path.parent.mkdir(parents=True, exist_ok=True)

    import matplotlib.pyplot as plt

    d = spec_df.head(top_k).copy()
    labels = d["spec_name"].astype(str) + " | lag=" + d["lag"].astype(int).astype(str)
    values = d["metric"].astype(float).values

    fig = plt.figure()
    ax = fig.add_subplot(111)
    ax.barh(range(len(values))[::-1], values[::-1])
    ax.set_yticks(range(len(values))[::-1])
    ax.set_yticklabels(labels[::-1])
    ax.set_xlabel("Effect magnitude |coef| (selected per spec)")
    ax.set_title("Top specification winners")
    fig.tight_layout()
    fig.savefig(out_path, dpi=200)
    plt.close(fig)
    return out_path
