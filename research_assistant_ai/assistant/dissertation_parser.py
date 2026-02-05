from __future__ import annotations
from dataclasses import dataclass
from pathlib import Path
from typing import List, Dict
import re

@dataclass
class DissertationSignals:
    title: str
    research_questions: List[str]
    keywords: List[str]

def _simple_keywords(text: str, top_k: int = 30) -> List[str]:
    text = re.sub(r"[^A-Za-z0-9\s]", " ", text).lower()
    stop = set([
        "the","and","of","to","in","a","for","is","that","with","as","on","are","be","or","by","an",
        "this","it","from","at","which","we","can","may","will","their","these","also","such",
    ])
    tokens = [t for t in text.split() if len(t) > 3 and t not in stop]
    freq: Dict[str, int] = {}
    for t in tokens:
        freq[t] = freq.get(t, 0) + 1
    return [w for w,_ in sorted(freq.items(), key=lambda x: x[1], reverse=True)[:top_k]]

def parse_docx(docx_path: Path) -> DissertationSignals:
    from docx import Document
    doc = Document(str(docx_path))
    paras = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    full = ""

    # include table text too
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                if cell.text:
                    full += "\n" + cell.text.strip()

    full = "\n".join(paras) + "\n" + full
    title = paras[0] if paras else docx_path.stem

    rqs: List[str] = []
    for p in paras:
        if re.search(r"\bRQ\b", p) or re.search(r"Research\s+Question", p, flags=re.I):
            rqs.append(p)

    if not rqs:
        for s in re.split(r"(?<=[\?])\s+", full):
            if "?" in s and len(s) < 260:
                rqs.append(s.strip())

    keywords = _simple_keywords(full, top_k=35)
    return DissertationSignals(title=title, research_questions=rqs[:10], keywords=keywords)
