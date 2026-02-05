from __future__ import annotations
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, Any, Optional, List
import json
import pandas as pd

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

@dataclass
class APATableSpec:
    table_number: int
    title: str
    notes: str

def _set_normal_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

def _add_apa_table(doc: Document, df: pd.DataFrame, spec: APATableSpec) -> None:
    # Table number (bold)
    p = doc.add_paragraph()
    r = p.add_run(f"Table {spec.table_number}")
    r.bold = True

    # Title (italic)
    p2 = doc.add_paragraph()
    r2 = p2.add_run(spec.title)
    r2.italic = True

    # Create table
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for j, col in enumerate(df.columns):
        hdr_cells[j].text = str(col)

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for j, col in enumerate(df.columns):
            val = row[col]
            cells[j].text = "" if pd.isna(val) else str(val)

    # Notes
    p3 = doc.add_paragraph()
    r3 = p3.add_run("Note. ")
    r3.italic = True
    p3.add_run(spec.notes)

def _fmt_p(p: float) -> str:
    try:
        pf = float(p)
    except Exception:
        return str(p)
    if pf < 0.001:
        return "< .001"
    return f"= {pf:.3f}"

def build_results_narrative(phase8_report: Dict[str, Any]) -> List[str]:
    """Create dissertation-ready narrative paragraphs from Phase 8 outputs.

    This intentionally avoids external factual claims. Where citations would typically be required
    (e.g., describing standard methods), the narrative inserts a placeholder token.
    """
    treated = phase8_report.get("treated_country", "RUS")
    interv = phase8_report.get("intervention_month", "")
    did = phase8_report.get("did", {})
    its = phase8_report.get("its", {})
    diag = phase8_report.get("diagnostics", {})
    sc = phase8_report.get("synthetic_control", {})

    pt = diag.get("parallel_trends_pretest", {})
    pre_joint = diag.get("event_study_preperiod_joint_test_fisher", {})

    paras: List[str] = []

    paras.append(
        f"This phase produced dissertation-oriented causal inference diagnostics and sensitivity analyses for the treated unit ({treated}) around the intervention month ({interv}). "
        f"The objective was to evaluate whether estimated post-intervention changes in cyber incidents remained directionally stable across identification strategies and robustness checks. "
        f"Method descriptions follow standard causal inference practice and should be accompanied by appropriate citations in the dissertation manuscript [CITATION NEEDED]."
    )

    paras.append(
        f"In the difference-in-differences model, the estimated average treatment effect on the treated was {did.get('att'):.4f} with p {_fmt_p(did.get('p_value'))}. "
        f"The parallel trends pre-test, implemented as a treated-by-trend interaction in the pre-intervention period, yielded a coefficient of {pt.get('coef_treated_x_trend'):.4f} with p {_fmt_p(pt.get('p_value'))}. "
        f"These results inform the plausibility of the identifying assumption that treated and control units followed comparable pre-intervention outcome trends [CITATION NEEDED]."
    )

    paras.append(
        f"The event study specification produced dynamic estimates of post-intervention changes relative to the omitted baseline month (k = -1). "
        f"A pre-period stability joint test using Fisher aggregation over pre-intervention coefficients produced p {_fmt_p(pre_joint.get('p_value'))}. "
        f"Together, the event study pattern and the pre-period stability assessment provide additional evidence regarding pre-intervention comparability and potential anticipation effects [CITATION NEEDED]."
    )

    if isinstance(sc, dict) and "base" in sc and "optimized" in sc:
        paras.append(
            f"Synthetic control estimation was conducted using a donor pool of control units, producing a baseline pre-intervention root mean squared error of {sc['base'].get('pre_rmse'):.4f}. "
            f"An optimized donor pool search reduced the pre-intervention root mean squared error to {sc['optimized'].get('pre_rmse'):.4f}. "
            f"These diagnostics are useful for assessing pre-intervention fit quality, which is a key requirement for credible synthetic control inference [CITATION NEEDED]."
        )

    paras.append(
        f"Interrupted time series estimation on the treated unit produced an estimated level change of {its.get('level_change'):.4f} with p {_fmt_p(its.get('p_level'))}, "
        f"and an estimated slope change of {its.get('slope_change'):.4f} with p {_fmt_p(its.get('p_slope'))}. "
        f"This specification supports evaluation of both immediate and gradual outcome changes following the intervention under a segmented regression framework [CITATION NEEDED]."
    )

    return paras

def export_results_docx(
    *,
    phase8_report: Dict[str, Any],
    event_study_df: pd.DataFrame,
    did_sensitivity_df: pd.DataFrame,
    out_path: Path
) -> Path:
    doc = Document()
    _set_normal_style(doc)

    # Title
    title = doc.add_paragraph("Phase 9: Dissertation-Ready Results Draft")
    title.runs[0].bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Narrative
    for para in build_results_narrative(phase8_report):
        doc.add_paragraph(para)

    # Tables
    _add_apa_table(
        doc,
        event_study_df,
        APATableSpec(
            table_number=1,
            title="Event Study Coefficients and p Values",
            notes="k denotes event time relative to the omitted baseline month (k = -1). p values are based on the event study model estimates. Replace [CITATION NEEDED] placeholders in the narrative with appropriate sources."
        )
    )

    _add_apa_table(
        doc,
        did_sensitivity_df,
        APATableSpec(
            table_number=2,
            title="Difference-in-Differences Sensitivity Across Covariate Sets",
            notes="Each row reports the estimated average treatment effect on the treated (ATT) and its p value for a specific covariate set. The covariate string is comma-delimited."
        )
    )

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path
